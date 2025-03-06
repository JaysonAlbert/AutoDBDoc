import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict, Any
from datetime import datetime
from .logger_config import logger

class DocGenerator:
    def __init__(self, db_reader, progress_callback=None):
        self.db_reader = db_reader
        self.doc = Document()
        self.progress_callback = progress_callback
        
    def generate_documentation(self, service_name, output_dir, selected_tables=None):
        """Generate documentation for the database.
        
        Args:
            service_name (str): Name of the database service
            output_dir (str): Directory to save the generated documentation
            selected_tables (list, optional): List of table names to document. If None, all tables will be documented.
        """
        try:
            if self.progress_callback:
                self.progress_callback('Starting documentation generation', 0, 100)

            # Get all tables or use selected tables
            if selected_tables is None:
                tables = self.db_reader.get_tables()
            else:
                # Validate that all selected tables exist
                all_tables = set(self.db_reader.get_tables())
                invalid_tables = set(selected_tables) - all_tables
                if invalid_tables:
                    raise ValueError(f"Invalid table names: {', '.join(invalid_tables)}")
                tables = selected_tables
            
            total_tables = len(tables)
            
            # Create title page
            self._create_title_page(service_name)
            if self.progress_callback:
                self.progress_callback('Created title page', 0, 100)

            # Create table index
            self._create_table_index(tables)
            if self.progress_callback:
                self.progress_callback('Created table index', 0, 100)

            
            # Generate documentation for each table
            for i, table in enumerate(tables, 1):
                if self.progress_callback:
                    self.progress_callback(f'Processing table {i}/{total_tables}: {table}', i, total_tables)
                self._document_table(table)
            
            # Save the document
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f'database_documentation_{timestamp}.docx'
            filepath = os.path.join(output_dir, filename)
            
            if self.progress_callback:
                self.progress_callback('Saving document...', 100, 100)
            
            self.doc.save(filepath)
            
            if self.progress_callback:
                self.progress_callback('Documentation completed', 100, 100)
            
            return filename
            
        except Exception as e:
            logger.error(f"Error generating documentation: {str(e)}")
            if self.progress_callback:
                self.progress_callback(f'Error: {str(e)}', 0, 100)
            raise
    
    def _create_title_page(self, service_name):
        """Create the title page of the document."""
        # Add title
        title = self.doc.add_heading('Database Documentation', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add service name
        service = self.doc.add_paragraph()
        service.alignment = WD_ALIGN_PARAGRAPH.CENTER
        service_run = service.add_run(service_name)
        service_run.font.size = Pt(16)
        
        # Add date
        date = self.doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date.add_run(datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        date_run.font.size = Pt(12)
        
    
    def _create_table_index(self, tables):
        """Create a table index page with links to each table section."""
        # Add heading
        self.doc.add_heading('Table Index', 1)
        
        # Add description
        description = self.doc.add_paragraph()
        description.add_run('This section provides a quick reference to all tables in the database with their descriptions.')
        
        # Create table with 2 columns: Table Name and Description
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        headers = ['Table Name', 'Description']
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add table rows
        for table_name in tables:
            # Get table description
            description = self.db_reader.get_table_description(table_name)
            
            # Add row
            row_cells = table.add_row().cells
            
            # Add table name with hyperlink
            paragraph = row_cells[0].paragraphs[0]
            run = paragraph.add_run(table_name)
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
            run.font.underline = True
            
            # Add description
            row_cells[1].text = description if description else 'No description available'
        
        # Add page break
        self.doc.add_page_break()
    
    def _document_table(self, table_name):
        """Document a single table."""
        try:
            # Get table information
            description = self.db_reader.get_table_description(table_name)
            columns = self.db_reader.get_table_columns(table_name)
            constraints = self.db_reader.get_table_constraints(table_name)
            column_comments = self.db_reader.get_column_comments(table_name)
            
            # Add table heading
            self.doc.add_heading(table_name, 1)
            
            # Add table description if available
            if description:
                self.doc.add_paragraph(description)
            
            # Add columns table
            self._add_columns_table(columns, column_comments)
            
            # Add constraints if any
            if constraints:
                self._add_constraints_section(constraints)
            
            # Add page break
            self.doc.add_page_break()
            
        except Exception as e:
            logger.error(f"Error documenting table {table_name}: {str(e)}")
            raise
    
    def _add_columns_table(self, columns, column_comments):
        """Add a table showing column information."""
        # Create table
        table = self.doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        headers = ['Name', 'Data Type', 'Length', 'Nullable', 'Default', 'Position', 'Description']
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for column in columns:
            row_cells = table.add_row().cells
            row_cells[0].text = column['name']
            row_cells[1].text = column['data_type']
            row_cells[2].text = str(column['length'])
            row_cells[3].text = 'Yes' if column['nullable'] == 'Y' else 'No'
            row_cells[4].text = str(column['default']) if column['default'] else ''
            row_cells[5].text = str(column['position'])
            row_cells[6].text = column_comments.get(column['name'], '')
            
            # Center align all cells
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def _add_constraints_section(self, constraints):
        """Add a section documenting table constraints."""
        # Add heading
        self.doc.add_heading('Constraints', 2)
        
        # Create table
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        headers = ['Name', 'Type', 'Details']
        
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for constraint in constraints:
            row_cells = table.add_row().cells
            row_cells[0].text = constraint['name']
            row_cells[1].text = constraint['type']
            
            # Format details based on constraint type
            if constraint['type'] == 'P':
                row_cells[2].text = f"Primary Key: {', '.join(constraint['columns'])}"
            elif constraint['type'] == 'R':
                row_cells[2].text = f"Foreign Key: {constraint['column']}"
            elif constraint['type'] == 'C':
                row_cells[2].text = f"Check: {constraint['condition']}"
            else:
                row_cells[2].text = str(constraint.get('condition', ''))
            
            # Center align all cells
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add spacing after table
        self.doc.add_paragraph() 